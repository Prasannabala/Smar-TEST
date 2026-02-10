"""
Document parser for extracting text from various file formats.
Supports TXT, PDF, DOCX files.
"""
import io
from pathlib import Path
from typing import Optional, Union, BinaryIO
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    """Represents a parsed document."""
    filename: str
    content: str
    file_type: str
    page_count: int = 1
    word_count: int = 0

    def __post_init__(self):
        self.word_count = len(self.content.split())


class DocumentParser:
    """
    Parser for extracting text from documents.
    Supports TXT, PDF, and DOCX formats.
    """

    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.doc'}

    @classmethod
    def parse(cls, file: Union[str, Path, BinaryIO], filename: Optional[str] = None) -> ParsedDocument:
        """
        Parse a document and extract its text content.

        Args:
            file: File path or file-like object
            filename: Optional filename (required if file is a file-like object)

        Returns:
            ParsedDocument with extracted content
        """
        if isinstance(file, (str, Path)):
            path = Path(file)
            filename = path.name
            extension = path.suffix.lower()

            with open(path, 'rb') as f:
                content, page_count = cls._extract_content(f, extension)
        else:
            if not filename:
                raise ValueError("filename is required when passing a file-like object")
            extension = Path(filename).suffix.lower()
            content, page_count = cls._extract_content(file, extension)

        return ParsedDocument(
            filename=filename,
            content=content,
            file_type=extension.lstrip('.'),
            page_count=page_count
        )

    @classmethod
    def _extract_content(cls, file: BinaryIO, extension: str) -> tuple[str, int]:
        """Extract content based on file extension."""
        if extension == '.txt':
            return cls._parse_txt(file), 1
        elif extension == '.pdf':
            return cls._parse_pdf(file)
        elif extension in ('.docx', '.doc'):
            return cls._parse_docx(file), 1
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    @staticmethod
    def _parse_txt(file: BinaryIO) -> str:
        """Parse a text file."""
        content = file.read()
        # Try different encodings
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        # Fallback: decode with errors ignored
        return content.decode('utf-8', errors='ignore')

    @staticmethod
    def _parse_pdf(file: BinaryIO) -> tuple[str, int]:
        """Parse a PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")

        reader = PdfReader(file)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

        return '\n\n'.join(pages), len(reader.pages)

    @staticmethod
    def _parse_docx(file: BinaryIO) -> str:
        """Parse a DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        doc = Document(file)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(' | '.join(row_text))

        return '\n\n'.join(paragraphs)

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if a file format is supported."""
        extension = Path(filename).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS

    @classmethod
    def get_supported_formats(cls) -> list[str]:
        """Get list of supported file formats."""
        return list(cls.SUPPORTED_EXTENSIONS)


class DocumentChunker:
    """
    Utility for chunking large documents to fit context windows.
    """

    @staticmethod
    def estimate_tokens(text: str) -> int:
        """Rough estimate of token count (approx 4 chars per token)."""
        return len(text) // 4

    @staticmethod
    def chunk_by_tokens(text: str, max_tokens: int = 4000, overlap: int = 200) -> list[str]:
        """
        Split text into chunks that fit within token limit.

        Args:
            text: Text to chunk
            max_tokens: Maximum tokens per chunk
            overlap: Token overlap between chunks

        Returns:
            List of text chunks
        """
        # Convert tokens to approximate character count
        max_chars = max_tokens * 4
        overlap_chars = overlap * 4

        if len(text) <= max_chars:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + max_chars

            # Try to break at paragraph or sentence boundary
            if end < len(text):
                # Look for paragraph break
                para_break = text.rfind('\n\n', start, end)
                if para_break > start + max_chars // 2:
                    end = para_break + 2
                else:
                    # Look for sentence break
                    for punct in ['. ', '.\n', '! ', '? ']:
                        sent_break = text.rfind(punct, start, end)
                        if sent_break > start + max_chars // 2:
                            end = sent_break + len(punct)
                            break

            chunks.append(text[start:end].strip())
            start = end - overlap_chars

        return chunks

    @staticmethod
    def summarize_sections(text: str, section_markers: list[str] = None) -> dict[str, str]:
        """
        Split document into sections based on markers.

        Args:
            text: Document text
            section_markers: List of section header patterns

        Returns:
            Dict mapping section names to content
        """
        if section_markers is None:
            section_markers = [
                'navigation', 'thumb rule', 'business rule',
                'requirement', 'constraint', 'overview',
                'introduction', 'summary', 'guideline'
            ]

        sections = {}
        lines = text.split('\n')
        current_section = "General"
        current_content = []

        for line in lines:
            line_lower = line.lower().strip()

            # Check if this line is a section header
            is_header = False
            for marker in section_markers:
                if marker in line_lower and len(line) < 100:
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    current_section = line.strip()
                    current_content = []
                    is_header = True
                    break

            if not is_header:
                current_content.append(line)

        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)

        return sections
